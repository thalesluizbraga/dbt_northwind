"""Gera diagrama de arquitetura e sumário executivo em Word com gráficos."""
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
CHARTS_DIR = ROOT / "5-bi_report" / "output"
ARCH_IMG = DOCS_DIR / "arquitetura_analytics_northwind.png"
DOC_PATH = ROOT / "Sumario_Executivo_Northwind.docx"

CHART_SECTIONS = [
    {
        "section": "3.1 Como aumentar o ticket médio?",
        "charts": [
            (
                "01_receita_e_qtd_pedidos_mensal.png",
                "Figura 1 — Evolução mensal da receita (base para monitoramento do ticket médio).",
            ),
            (
                "03_impacto_desconto.png",
                "Figura 2 — Impacto do desconto: ticket médio com vs sem desconto.",
            ),
            (
                "04_cross_sell_categorias.png",
                "Figura 3 — Top pares de categorias compradas juntas (oportunidades de cross-sell).",
            ),
            (
                "05_produtos_ancora_expansao.png",
                "Figura 4 — Produtos âncora (1º pedido) vs produtos de expansão (pedidos seguintes).",
            ),
        ],
    },
    {
        "section": "3.2 Como reduzir o churn?",
        "charts": [
            (
                "08_cohort_retencao.png",
                "Figura 5 — Coorte de retenção: heatmap por mês da coorte e curva média ponderada.",
            ),
            (
                "09_clientes_em_declinio.png",
                "Figura 6 — Clientes em declínio: periodicidade média vs dias desde o último pedido.",
            ),
            (
                "10_perfil_churn_vs_ativos.png",
                "Figura 7 — Perfil geográfico: potencial churn vs clientes ativos (país, região, cidade).",
            ),
        ],
    },
    {
        "section": "3.3 Onde estão os melhores performers?",
        "charts": [
            (
                "01_ticket_medio_mensal.png",
                "Figura 8 — Evolução do ticket médio mensal.",
            ),
            (
                "02_itens_unidades_por_pedido.png",
                "Figura 9 — Itens e unidades médias por pedido (composição do ticket).",
            ),
        ],
    },
]


def _rounded_box(ax, xy, w, h, color, text, fontsize=9, text_color="#1e293b", edge="#94a3b8", linestyle="-"):
    box = FancyBboxPatch(
        xy,
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.5,
        edgecolor=edge,
        facecolor=color,
        linestyle=linestyle,
    )
    ax.add_patch(box)
    ax.text(xy[0] + w / 2, xy[1] + h / 2, text, ha="center", va="center", fontsize=fontsize, color=text_color, weight="bold")
    return box


def _arrow(ax, start, end):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.8,
            color="#334155",
            shrinkA=2,
            shrinkB=2,
        )
    )


def build_architecture_diagram() -> Path:
    fig, ax = plt.subplots(figsize=(16, 8))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor("#f8fafc")

    outer = FancyBboxPatch(
        (0.2, 0.3),
        15.6,
        7.4,
        boxstyle="round,pad=0.02,rounding_size=0.15",
        linewidth=2,
        edgecolor="#cbd5e1",
        facecolor="#f1f5f9",
    )
    ax.add_patch(outer)

    ax.text(0.6, 7.35, "Arquitetura da Solução — Northwind Traders", fontsize=16, weight="bold", color="#0f172a")
    ax.text(0.6, 6.95, "Pipeline de Analytics Engineering", fontsize=11, color="#475569")

    infra_colors = [("#dbeafe", "Docker"), ("#bfdbfe", "Linux"), ("#93c5fd", "Local\nSetup")]
    for i, (color, label) in enumerate(infra_colors):
        _rounded_box(ax, (0.55, 2.0 + i * 1.35), 1.35, 1.0, color, label, fontsize=10)

    orch = FancyBboxPatch(
        (2.3, 1.5),
        10.0,
        5.2,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=2.2,
        edgecolor="#dc2626",
        facecolor="none",
        linestyle=(0, (6, 4)),
    )
    ax.add_patch(orch)
    ax.text(7.3, 6.35, "ORQUESTRAÇÃO", ha="center", fontsize=12, weight="bold", color="#dc2626")

    steps = [
        (2.7, 4.7, 1.7, 1.2, "#ffffff", "DATA\nSOURCE", "CSVs\n1-data/"),
        (4.7, 4.7, 1.8, 1.2, "#ffffff", "EXTRACT\n& LOAD", "Python\nETL"),
        (6.8, 4.7, 1.9, 1.2, "#ffffff", "DATALAKE", "PostgreSQL\nschema raw"),
        (9.0, 4.7, 1.7, 1.2, "#ffffff", "TRANSFORM", "dbt\nstaging/marts"),
        (11.1, 4.7, 1.9, 1.2, "#ffffff", "DATA\nWAREHOUSE", "PostgreSQL\nfct_* / dim_*"),
    ]
    for x, y, w, h, color, title, subtitle in steps:
        _rounded_box(ax, (x, y), w, h, color, f"{title}\n{subtitle}", fontsize=8)

    for x1, x2 in [(4.4, 4.7), (6.5, 6.8), (8.7, 9.0), (10.8, 11.1)]:
        _arrow(ax, (x1, 5.3), (x2, 5.3))

    _rounded_box(ax, (5.8, 2.0), 3.0, 1.1, "#fff7ed", "Apache Airflow\n(Cosmos / DAG)", fontsize=9, edge="#ea580c")
    ax.text(7.3, 1.75, "Agendamento e orquestração do pipeline dbt", ha="center", fontsize=8, color="#9a3412")

    viz_box = FancyBboxPatch(
        (12.6, 1.5),
        2.9,
        5.2,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=1.8,
        edgecolor="#2563eb",
        facecolor="#eff6ff",
    )
    ax.add_patch(viz_box)
    ax.text(14.05, 6.35, "VISUALIZAÇÃO\nDE DADOS", ha="center", fontsize=11, weight="bold", color="#1d4ed8")

    viz_tools = ["Jupyter\nNotebook", "Power BI", "Sumário\nExecutivo"]
    for i, tool in enumerate(viz_tools):
        _rounded_box(ax, (12.85, 4.8 - i * 1.35), 2.4, 1.0, "#ffffff", tool, fontsize=9, edge="#60a5fa")

    _arrow(ax, (13.0, 5.3), (12.85, 5.3))
    ax.annotate(
        "",
        xy=(12.6, 5.3),
        xytext=(13.0, 5.3),
        arrowprops=dict(arrowstyle="-|>", color="#334155", lw=1.8),
    )

    ax.text(14.05, 0.75, "dbt_northwind | Analytics Engineering", ha="center", fontsize=9, color="#64748b")

    fig.tight_layout()
    DOCS_DIR.mkdir(exist_ok=True)
    fig.savefig(ARCH_IMG, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return ARCH_IMG


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 58, 138)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_chart(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[Gráfico não encontrado: {image_path.name}]", bold=True)
        return
    doc.add_paragraph()
    doc.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(71, 85, 105)
    doc.add_paragraph()


def build_document(arch_img: Path) -> Path:
    doc = Document()

    title = doc.add_heading("Sumário Executivo — Northwind Traders", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Relatório de Indicadores de Performance | Analytics Engineering")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
        run.italic = True
    doc.add_paragraph()

    add_heading(doc, "1. Objetivo")
    add_paragraph(
        doc,
        "A Northwind Traders opera com dados dispersos em planilhas sob demanda, gerando inconsistências "
        "entre áreas e dificultando decisões estratégicas. Com cerca de 30 funcionários e faturamento mensal "
        "estimado em R$ 1,5 milhão, a empresa busca uma visão integrada dos dados do ERP para sustentar "
        "o crescimento acelerado.",
    )
    add_paragraph(doc, "As perguntas de negócio centrais do desafio são:", bold=True)
    add_bullets(
        doc,
        [
            "Como aumentar o ticket médio por pedido e por cliente?",
            "Como reduzir o churn (perda ou não renovação de clientes B2B)?",
            "Quais regiões, categorias e clientes concentram a melhor performance de receita?",
            "Quais ações práticas podem ser tomadas com base nos indicadores para melhorar os resultados?",
        ],
    )
    add_paragraph(
        doc,
        "Este sumário consolida as respostas obtidas a partir do pipeline de analytics engineering "
        "implementado no projeto dbt_northwind e das análises documentadas em northwind_analyses.ipynb.",
    )

    add_heading(doc, "2. Arquitetura de Analytics Engineering")
    add_paragraph(
        doc,
        "A solução segue o padrão moderno de analytics engineering, com camadas bem definidas: "
        "ingestão (Python), armazenamento (PostgreSQL), transformação (dbt), orquestração (Airflow) "
        "e visualização (Jupyter, Power BI e este relatório). O diagrama abaixo espelha o fluxo "
        "Extract → Load → Transform → Warehouse → Visualização.",
    )
    add_bullets(
        doc,
        [
            "Infraestrutura: Docker + Linux no ambiente local (2-local_setup).",
            "Data Source: 14 tabelas CSV do ERP Northwind (1-data/).",
            "Extract & Load: Python (load_raw.py) → PostgreSQL schema raw.",
            "Transform: dbt (staging → intermediate → marts: dim_*, fct_*, agg_*).",
            "Data Warehouse: PostgreSQL com modelos analíticos materializados.",
            "Orquestração: Apache Airflow + Astronomer Cosmos (DAG northwind_pipeline).",
            "Visualização: Jupyter (northwind_analyses.ipynb), Power BI e sumário executivo.",
        ],
    )
    doc.add_paragraph()
    doc.add_picture(str(arch_img), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figura A — Arquitetura da solução de analytics engineering implementada.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True

    add_heading(doc, "3. Respostas às Perguntas de Negócio")

    add_heading(doc, "3.1 Como aumentar o ticket médio?", level=2)
    add_paragraph(doc, "Principais achados das análises:", bold=True)
    add_bullets(
        doc,
        [
            "Concentração de receita: os 10 maiores clientes respondem por 45% da receita total da carteira, "
            "com ticket médio entre R$ 1.556 (Folk och fä HB) e R$ 3.938 (QUICK-Stop).",
            "Combinações de alto valor: Beverages + Dairy Products (ticket médio R$ 2.444) e "
            "Beverages + Condiments (R$ 2.414) no mesmo pedido.",
            "Oportunidades regionais: Meat/Poultry em Co. Cork (R$ 4.183) e Beverages em Québec (R$ 3.139).",
            "Produtos âncora atraem novos clientes; produtos de expansão elevam valor em pedidos subsequentes.",
            "Pedidos com desconto (45,8%) têm ticket médio de R$ 1.702 vs R$ 1.375 sem desconto.",
        ],
    )
    for fname, caption in CHART_SECTIONS[0]["charts"]:
        add_chart(doc, CHARTS_DIR / fname, caption)

    add_heading(doc, "3.2 Como reduzir o churn?", level=2)
    add_paragraph(doc, "Principais achados das análises:", bold=True)
    add_bullets(
        doc,
        [
            "Coorte de retenção: queda de 100% para ~39% no mês 1 após o primeiro pedido.",
            "Periodicidade média global: 83 dias entre compras; 17 clientes (19%) em potencial churn.",
            "Clientes críticos: Centro comercial Moctezuma (657 dias), Mère Paillarde (top 10 receita + risco).",
            "Perfil do churn: França (29%) e EUA (24%) concentram potenciais churns.",
        ],
    )
    for fname, caption in CHART_SECTIONS[1]["charts"]:
        add_chart(doc, CHARTS_DIR / fname, caption)

    add_heading(doc, "3.3 Onde estão os melhores performers?", level=2)
    add_bullets(
        doc,
        [
            "Clientes: QUICK-Stop, Ernst Handel e Save-a-lot Markets lideram receita e frequência.",
            "Região × categoria: Meat/Poultry e Beverages no top 10 de ticket médio regional.",
            "Cross-sell: Beverages é central nas combinações de maior ticket.",
        ],
    )
    for fname, caption in CHART_SECTIONS[2]["charts"]:
        add_chart(doc, CHARTS_DIR / fname, caption)

    add_heading(doc, "3.4 Visão integrada dos dados", level=2)
    add_paragraph(
        doc,
        "O pipeline centraliza dados brutos, modelos dbt testados (49 testes aprovados) e indicadores "
        "reutilizáveis (fct_order_items, dim_customers, dim_customer_status, agregados mensais). "
        "Os gráficos acima foram gerados automaticamente a partir do notebook northwind_analyses.ipynb.",
    )

    add_heading(doc, "4. Conclusão e Ações Recomendadas")
    add_paragraph(
        doc,
        "Com base nas análises e visualizações apresentadas, recomenda-se:",
    )

    add_heading(doc, "4.1 Aumento de ticket médio", level=2)
    add_bullets(
        doc,
        [
            "Campanhas de cross-sell com pares Beverages + Dairy/Condiments por região de alto ticket.",
            "Playbook comercial dos top 10 clientes (45% da receita).",
            "Produtos âncora na aquisição; produtos de expansão no upsell pós-primeiro pedido.",
            "Revisar descontos: manter em pedidos de alto volume, evitar em pedidos abaixo de R$ 1.375.",
        ],
    )

    add_heading(doc, "4.2 Redução de churn", level=2)
    add_bullets(
        doc,
        [
            "Programa de retenção no mês 1 pós-primeiro pedido.",
            "Monitoramento semanal dos 17 clientes em potencial churn.",
            "Reativação prioritária na França e EUA.",
            "Ação comercial imediata para Mère Paillarde e demais top clientes em risco.",
        ],
    )

    add_heading(doc, "4.3 Governança e expansão do BI", level=2)
    add_bullets(
        doc,
        [
            "Dashboards em Power BI conectados às tabelas marts.",
            "Pipeline dbt + Airflow para atualização recorrente.",
            "Integração futura de CRM (Salesforce) e ContaAzul no mesmo padrão de camadas.",
        ],
    )

    add_paragraph(
        doc,
        "A arquitetura implementada oferece base escalável para decisões data driven, com indicadores "
        "visuais que respondem diretamente às perguntas estratégicas da Northwind.",
    )

    doc.save(DOC_PATH)
    return DOC_PATH


if __name__ == "__main__":
    img = build_architecture_diagram()
    path = build_document(img)
    print(f"Diagrama: {img}")
    print(f"Documento: {path}")
    missing = []
    for section in CHART_SECTIONS:
        for fname, _ in section["charts"]:
            if not (CHARTS_DIR / fname).exists():
                missing.append(fname)
    if missing:
        print("Gráficos ausentes:", ", ".join(missing))
    else:
        print(f"Gráficos incluídos: {sum(len(s['charts']) for s in CHART_SECTIONS)}")
